"""Two-column TECHNICAL SKILLS comes from the RENDERER, never the face.

The face keeps pipe-free "Category: a, b, c" lines, so the format gate (which
rejects | off the contact line) and the no-drift gate read the same flat text
they always did; only the layout gains the second column.
"""

from __future__ import annotations

from pathlib import Path

import job_agent.tailor as tailor_pkg
from job_agent.tailor.career_facts import load_career_facts
from job_agent.tailor.render_pdf import render_docx, render_pdf
from job_agent.tailor.tailor import tailor_resume
from job_agent.tailor.verify import extract_pdf_text, verify_format

DEMO = Path(tailor_pkg.__file__).parent / "demo"
FACTS = load_career_facts(DEMO / "demo_career_facts.yaml")
RESULT = tailor_resume(FACTS, (DEMO / "demo_jd.txt").read_text(),
                       stub_response=(DEMO / "demo_response.txt").read_text(),
                       megaprompt="(x)")
# The demo face's three skill categories, exactly as they appear on the face.
CATEGORIES = ("Languages", "Data & Orchestration", "Cloud")


def test_face_stays_pipe_free_and_passes_the_format_gate():
    for line in RESULT.resume_text.splitlines():
        if any(line.startswith(c) for c in CATEGORIES):
            assert "|" not in line
    verify_format(RESULT, FACTS)   # gate unchanged, face unchanged


def test_pdf_skills_render_in_two_columns(tmp_path):
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer

    pdf = render_pdf(RESULT.resume_text, tmp_path / "resume.pdf")
    xs: dict[str, float] = {}
    for page in extract_pages(str(pdf)):
        for el in page:
            if isinstance(el, LTTextContainer):
                for line in el:
                    text = line.get_text().strip()
                    for cat in CATEGORIES:
                        if text.startswith(f"{cat}:"):
                            xs.setdefault(cat, line.x0)
    assert set(xs) == set(CATEGORIES), f"categories missing from PDF: {xs}"
    # Row-major two-column layout: category #2 sits in the RIGHT column,
    # #1 and #3 in the left — two distinct x origins, far apart.
    assert xs["Data & Orchestration"] - xs["Languages"] > 100
    assert abs(xs["Cloud"] - xs["Languages"]) < 5


def test_pdf_two_column_skills_still_extract_completely(tmp_path):
    pdf = render_pdf(RESULT.resume_text, tmp_path / "resume.pdf")
    text = extract_pdf_text(pdf)
    for needle in ("Languages", "Python, SQL", "Apache Airflow", "AWS (S3, Glue, Redshift)"):
        assert needle in text


def test_docx_skills_render_as_a_two_column_table(tmp_path):
    from docx import Document

    path = render_docx(RESULT.resume_text, tmp_path / "resume.docx")
    doc = Document(str(path))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.columns) == 2
    cell_text = " | ".join(c.text for row in table.rows for c in row.cells)
    for cat in CATEGORIES:
        assert f"{cat}:" in cell_text
    # the categories moved INTO the table — no stray flat paragraphs remain
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Languages:" not in body_text
