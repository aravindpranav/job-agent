"""Render tailored resume text into a clean, ATS-safe PDF and matching .docx.

Format (mirrors a clean single-column reference resume):
- Header: Name on line 1; "email | phone" on line 2. No tagline.
- CAPS section headings (PROFESSIONAL SUMMARY, …) each under a thin rule.
- Technical Skills: the face stays pipe-free "Category: value, value" lines;
  the LAYOUT renders them as a two-column borderless table (structure from
  layout, so the format gate's no-pipe rule still holds on the text).
- Experience blocks: Role/Company/Project Description/Duration, then bullets.
- Real "•" bullets that EXTRACT as text — achieved with an embedded Unicode font
  (DejaVu Sans, bundled) since reportlab's base-14 Helvetica can't map "•" for
  extraction. No em-dashes, no separator lines, black text on white.

The PDF's text is verified (verify.py) to be selectable with sections in order.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from xml.sax.saxutils import escape

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from job_agent.tailor.textnorm import strip_markdown

_ASSETS = Path(__file__).resolve().parent / "assets"
FONT, FONT_BOLD = "DejaVuSans", "DejaVuSans-Bold"
_FONTS_READY = False


def _ensure_fonts() -> None:
    global _FONTS_READY
    if not _FONTS_READY:
        pdfmetrics.registerFont(TTFont(FONT, str(_ASSETS / "DejaVuSans.ttf")))
        pdfmetrics.registerFont(TTFont(FONT_BOLD, str(_ASSETS / "DejaVuSans-Bold.ttf")))
        # Register as a family so inline <b>…</b> maps to the embedded bold face
        # (renders bold AND still extracts as text).
        pdfmetrics.registerFontFamily(FONT, normal=FONT, bold=FONT_BOLD,
                                      italic=FONT, boldItalic=FONT_BOLD)
        _FONTS_READY = True


# Sub-labels whose VALUE should be bold (company name, role title).
_BOLD_VALUE_LABELS = {"role", "company"}


CANONICAL_HEADINGS = [
    "Professional Summary",
    "Technical Skills",
    "Professional Experience",
    "Education",
    "Certifications",
]
_HEADING_LOOKUP = {h.lower(): h for h in CANONICAL_HEADINGS}
_SUBLABELS = re.compile(
    r"^(Role|Company|Project Description|Duration|Responsibilities|Achievements):\s*(.*)$",
    re.IGNORECASE,
)
_CATEGORY = re.compile(r"^([A-Za-z][A-Za-z0-9 &/+\-]{1,44}):\s+(\S.*)$")
_SEPARATOR = re.compile(r"^[\-_.=|*·•\s]+$")  # a line that is only rule/separator chars


def _canonical_heading(line: str) -> str | None:
    return _HEADING_LOOKUP.get(line.strip().rstrip(":").strip().lower())


def clean_resume_text(text: str) -> str:
    """Remove em-dashes and separator lines, and tighten blank-line padding the
    model adds between bullets (so the résumé doesn't run long)."""
    text = text.replace(" — ", ", ").replace("—", ", ").replace("−", "-")
    text = re.sub(r"\s–\s", " - ", text).replace("–", "-")
    lines = [ln for ln in text.splitlines() if not _SEPARATOR.fullmatch(ln.strip())]

    tight: list[str] = []
    for i, ln in enumerate(lines):
        if not ln.strip():
            prev = tight[-1].strip() if tight else ""
            nxt = next((s.strip() for s in lines[i + 1:] if s.strip()), "")
            if not prev:
                continue                                   # collapse consecutive/leading blanks
            if prev[:1] in "-•*" or nxt[:1] in "-•*":
                continue                                   # no blank line around bullets
        tight.append(ln)

    out = "\n".join(tight)
    out = re.sub(r",\s*,", ",", out)
    return re.sub(r"[ \t]{2,}", " ", out)


# Page budget for the fit loop: four full employers + grouped skills need three
# pages; trimming beyond that would push roles back to skeletal density.
MAX_RESUME_PAGES = 3

# Never trim a role below this many responsibility bullets when fitting pages.
# Matches the prompt's older-role depth floor (3-4): the fit loop may remove
# surplus, never hollow a role back to 2-bullet density.
_RESP_FLOOR = 3


def drop_last_responsibility(resume_text: str) -> str:
    """Drop one responsibility bullet — the last (least JD-relevant) one of the
    role that currently has the most — without touching achievements. Used to fit
    the résumé to the page budget. Returns the text unchanged if nothing is
    droppable."""
    lines = resume_text.splitlines()
    roles: list[list[int]] = []
    role_idx, section = -1, None
    for i, raw in enumerate(lines):
        line = strip_markdown(raw)
        if re.match(r"(?i)^Role:", line):
            role_idx += 1
            roles.append([])
            section = None
        elif re.match(r"(?i)^Responsibilities:", line):
            section = "resp"
        elif re.match(r"(?i)^Achievements:", line) or _canonical_heading(line):
            section = None
        elif section == "resp" and raw.strip()[:1] in "-•*":
            roles[role_idx].append(i)

    candidates = [(len(idx), ri) for ri, idx in enumerate(roles) if len(idx) > _RESP_FLOOR]
    if not candidates:
        return resume_text
    _, ri = max(candidates)
    drop = roles[ri][-1]
    return "\n".join(ln for j, ln in enumerate(lines) if j != drop)


def trim_to_caps(resume_text: str) -> str:
    """Enforce per-role bullet caps by keeping the first N bullets of each section.

    The model is instructed to rank bullets strongest-first, so keeping the first N
    keeps the most JD-relevant ones. Caps: most-recent role 6 responsibilities,
    older roles 4; achievements 3 per role. Only drops bullets — never edits or
    reorders — so it can't affect the no-drift or banked-metric checks beyond
    removing surplus lines.
    """
    role_idx = -1
    section: str | None = None
    kept = 0
    cap = 0
    out: list[str] = []
    for raw in resume_text.splitlines():
        line = strip_markdown(raw)
        if re.match(r"(?i)^Role:", line):
            role_idx += 1
            section = None
        elif re.match(r"(?i)^Responsibilities:", line):
            section, kept, cap = "resp", 0, (6 if role_idx == 0 else 4)
        elif re.match(r"(?i)^Achievements:", line):
            section, kept, cap = "ach", 0, 3
        elif _canonical_heading(line):
            section = None
        elif section and raw.strip()[:1] in "-•*":
            if kept >= cap:
                continue  # drop surplus bullet
            kept += 1
        out.append(raw)
    return "\n".join(out)


def normalize_header(resume_text: str, name: str, email: str, phone: str) -> str:
    """Force the header to 'Name' / 'email | phone', dropping any model tagline."""
    lines = resume_text.splitlines()
    idx = next((i for i, ln in enumerate(lines)
                if _canonical_heading(strip_markdown(ln))), 0)
    body = "\n".join(lines[idx:]).lstrip("\n")
    return f"{name}\n{email} | {phone}\n\n{body}"


SKILL_FONT_SIZE = 9.5


def _styles() -> dict[str, ParagraphStyle]:
    base = ParagraphStyle("body", fontName=FONT, fontSize=9.5, leading=12,
                          leftIndent=0, firstLineIndent=0, alignment=0)
    return {
        # Name and contact share identical geometry (leftIndent/firstLineIndent/
        # alignment) so there is zero horizontal offset between them.
        "name": ParagraphStyle("name", parent=base, fontName=FONT_BOLD, fontSize=15,
                               leftIndent=0, firstLineIndent=0, alignment=0, spaceAfter=1),
        "contact": ParagraphStyle("contact", parent=base, fontSize=9.5,
                                  leftIndent=0, firstLineIndent=0, alignment=0, spaceAfter=4),
        "heading": ParagraphStyle("heading", parent=base, fontName=FONT_BOLD, fontSize=11,
                                  spaceBefore=6, spaceAfter=1),
        "body": base,
        "bullet": ParagraphStyle("bullet", parent=base, leftIndent=13, firstLineIndent=-9,
                                 spaceAfter=1),
    }


# A modest, fixed hanging indent: wrapped skill values align in a consistent
# indented column instead of returning to the margin. Deliberately small — a
# large (label-width) hang makes PDF text extractors read the wrapped lines out
# of order, scrambling the skills for an ATS. This keeps the linear reading order.
SKILL_HANG = 14


def _skill_style() -> ParagraphStyle:
    return ParagraphStyle("skill", fontName=FONT, fontSize=SKILL_FONT_SIZE, leading=12,
                          leftIndent=SKILL_HANG, firstLineIndent=-SKILL_HANG, spaceAfter=2.5)


def _parse_lines(resume_text: str):
    """Yield (kind, payload) describing each line, tracking the current section."""
    meaningful = 0
    section = None
    for raw in resume_text.splitlines():
        line = strip_markdown(raw)
        if not line:
            yield ("blank", "")
            continue
        meaningful += 1
        if meaningful == 1:
            yield ("name", line)
        elif meaningful == 2:
            yield ("contact", line)
        elif (heading := _canonical_heading(line)):
            section = heading
            yield ("heading", heading.upper())
        elif line[0] in "-•*":
            yield ("bullet", line[1:].strip())
        elif (m := _SUBLABELS.match(line)):
            yield ("label", (m.group(1), m.group(2)))
        elif section == "Technical Skills" and (m := _CATEGORY.match(line)):
            yield ("category", (m.group(1), m.group(2)))
        else:
            yield ("body", line)


# Two columns across the usable letter width (8.5" minus 0.6" margins each side).
_SKILL_COL_WIDTH = 3.65 * inch


def _skills_table(categories: list[tuple[str, str]]) -> Table:
    """Two-column, borderless layout for the buffered skill-category lines.

    Row-major fill (1|2, 3|4, …) so an ATS reading left-to-right per row keeps
    a sensible order; each cell is a normal Paragraph, so the text extracts
    exactly like the old single-column lines did.
    """
    style = _skill_style()
    cells = [Paragraph(f"<b>{escape(label)}:</b> {escape(values)}", style)
             for label, values in categories]
    rows = [cells[i:i + 2] for i in range(0, len(cells), 2)]
    if len(rows[-1]) == 1:
        rows[-1].append("")
    table = Table(rows, colWidths=[_SKILL_COL_WIDTH, _SKILL_COL_WIDTH], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
    ]))
    return table


def render_pdf(resume_text: str, out_path: str | Path) -> Path:
    """Write the ATS-safe PDF (single column; skills in a two-column table).
    Returns the path."""
    _ensure_fonts()
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    styles = _styles()
    flow = []
    skill_cats: list[tuple[str, str]] = []

    def flush_skills() -> None:
        if skill_cats:
            flow.append(_skills_table(skill_cats))
            skill_cats.clear()

    for kind, payload in _parse_lines(resume_text):
        if kind == "category":
            skill_cats.append(payload)
            continue
        if kind == "blank" and skill_cats:
            continue        # keep buffering across blank lines between categories
        flush_skills()
        if kind == "blank":
            flow.append(Spacer(1, 4))
        elif kind == "name":
            flow.append(Paragraph(escape(payload), styles["name"]))
        elif kind == "contact":
            flow.append(Paragraph(escape(payload), styles["contact"]))
        elif kind == "heading":
            flow.append(Paragraph(escape(payload), styles["heading"]))
            flow.append(HRFlowable(width="100%", thickness=0.6, color=colors.grey,
                                   spaceBefore=1, spaceAfter=4))
        elif kind == "bullet":
            flow.append(Paragraph("• " + escape(payload), styles["bullet"]))
        elif kind == "label":
            label, rest = payload
            if label.lower() in _BOLD_VALUE_LABELS:  # bold company name / role title
                flow.append(Paragraph(f"<b>{escape(label)}: {escape(rest)}</b>", styles["body"]))
            else:
                flow.append(Paragraph(f"<b>{escape(label)}:</b> {escape(rest)}", styles["body"]))
        else:
            flow.append(Paragraph(escape(payload), styles["body"]))
    flush_skills()   # skills were the last section on the face

    SimpleDocTemplate(
        str(out_path), pagesize=letter,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch, title="Resume",
    ).build(flow)
    return out_path


def render_docx(resume_text: str, out_path: str | Path) -> Path:
    """Write a professional single-column .docx.

    Layout: large bold name with the contact line beneath; CAPS section headings
    under a thin rule; each experience block shows the bold Company with its dates
    right-aligned on the same line, the role title in italics below, then the
    project description and bulleted responsibilities/achievements. Clean Calibri
    font, comfortable spacing. This .docx is the source of truth — the PDF is
    produced from it via LibreOffice so the two match exactly.
    """
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Emu, Inches, Pt

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.0

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.5)
    right_edge = Emu(sec.page_width - sec.left_margin - sec.right_margin)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr()          # thin rule = bottom border
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "2"), ("w:color", "999999")):
            bottom.set(qn(k), v)
        borders.append(bottom)
        pPr.append(borders)

    pending: dict[str, str] = {}
    skill_cats: list[tuple[str, str]] = []

    def flush_skills() -> None:
        # Two-column, borderless table (row-major), mirroring the PDF layout.
        if not skill_cats:
            return
        table = doc.add_table(rows=(len(skill_cats) + 1) // 2, cols=2)
        for i, (label, rest) in enumerate(skill_cats):
            cell = table.cell(i // 2, i % 2)
            cell.width = Inches(3.65)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"{label}: ").bold = True
            p.add_run(rest)
        skill_cats.clear()

    def flush_experience_header(duration: str) -> None:
        # Company (bold) + dates right-aligned on the same line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
        p.add_run(pending.get("company", "")).bold = True
        p.add_run("\t" + duration)
        # Role title in italics beneath the company
        if pending.get("role"):
            pr = doc.add_paragraph()
            pr.paragraph_format.space_after = Pt(1)
            pr.add_run(pending["role"]).italic = True
        # Project description line (kept, project-focused)
        if pending.get("projdesc"):
            pd = doc.add_paragraph()
            pd.add_run("Project Description: ").bold = True
            pd.add_run(pending["projdesc"])
        pending.clear()

    for kind, payload in _parse_lines(resume_text):
        if kind == "category":
            skill_cats.append(payload)
            continue
        if kind == "blank":
            continue
        flush_skills()
        if kind == "name":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(payload); r.bold = True; r.font.size = Pt(18)
        elif kind == "contact":
            p = doc.add_paragraph(payload)
            p.paragraph_format.space_after = Pt(2)
        elif kind == "heading":
            heading(payload)
        elif kind == "bullet":
            b = doc.add_paragraph(f"•  {payload}")
            b.paragraph_format.left_indent = Inches(0.2)
            b.paragraph_format.first_line_indent = Inches(-0.16)
        elif kind == "label":
            label, rest = payload
            key = label.lower()
            if key in ("role", "company", "project description"):
                pending["projdesc" if key == "project description" else key] = rest
            elif key == "duration":
                flush_experience_header(rest)
            else:  # Responsibilities: / Achievements:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.add_run(f"{label}:").bold = True
        else:
            doc.add_paragraph(payload)
    flush_skills()   # skills were the last section on the face

    doc.save(str(out_path))
    return out_path


def find_soffice() -> str | None:
    """Locate the LibreOffice/soffice binary, or None if not installed."""
    for name in ("soffice", "libreoffice"):
        if path := shutil.which(name):
            return path
    for path in ("/Applications/LibreOffice.app/Contents/MacOS/soffice",
                 str(Path.home() / "Applications/LibreOffice.app/Contents/MacOS/soffice")):
        if Path(path).exists():
            return path
    return None


def docx_to_pdf(docx_path: str | Path, out_dir: str | Path) -> Path | None:
    """Convert a .docx to PDF with LibreOffice so the PDF matches the .docx
    exactly. Returns the PDF path, or None if LibreOffice isn't available.

    Deletes the target first and confirms a fresh write — repeated soffice calls
    otherwise delegate to a running instance and can leave a stale PDF behind.
    """
    soffice = find_soffice()
    if not soffice:
        return None
    docx_path, out_dir = Path(docx_path), Path(out_dir)
    pdf = out_dir / f"{docx_path.stem}.pdf"
    pdf.unlink(missing_ok=True)
    profile = f"file://{tempfile.gettempdir()}/jobagent_libreoffice_profile"
    for _ in range(3):
        subprocess.run(
            [soffice, f"-env:UserInstallation={profile}", "--headless",
             "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
            capture_output=True, timeout=180, check=False,
        )
        if pdf.exists() and pdf.stat().st_size > 0:
            return pdf
    return pdf if pdf.exists() else None
